"""DOCX document parser using python-docx."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from app.services.doc_parser.base import BaseDocParser, ParsedDocument

_MONOSPACE_FONTS = {"courier new", "consolas", "monaco", "menlo", "source code pro", "fira code", "jetbrains mono"}


class DocxParser(BaseDocParser):
    name = "docx"

    def can_parse(self, mime_type: str, filename: str) -> bool:
        return mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ) or filename.endswith((".docx", ".doc"))

    def _parse_sync(self, content: bytes, filename: str) -> ParsedDocument:
        doc = Document(BytesIO(content))

        headings: list[dict] = []
        paragraphs: list[str] = []
        code_blocks: list[dict] = []
        pos = 0

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            level = self._guess_heading_level(style_name)
            if level > 0:
                headings.append({"level": level, "text": text, "position": pos})
                paragraphs.append(text)
                pos += len(text) + 1
            elif self._is_code_block(para):
                code_blocks.append({"language": "", "code": text, "position": pos})
                paragraphs.append(text)
                pos += len(text) + 1
            else:
                paragraphs.append(text)
                pos += len(text) + 1

        tables = self._extract_tables(doc.tables)

        return ParsedDocument(
            text="\n\n".join(paragraphs),
            headings=headings,
            code_blocks=code_blocks,
            tables=tables,
            images=[],
            metadata={
                "title": filename,
                "author": (doc.core_properties.author or ""),
                "page_count": len(paragraphs),
            },
        )

    @staticmethod
    def _guess_heading_level(style_name: str) -> int:
        style_lower = style_name.lower()
        if "heading" in style_lower or "heading" in style_name:
            for i in range(9, 0, -1):
                if str(i) in style_name:
                    return i
            return 1
        return 0

    @staticmethod
    def _is_code_block(para) -> bool:
        font_name = ""
        for run in para.runs:
            if run.font.name:
                font_name = run.font.name.lower()
                break
        if font_name in _MONOSPACE_FONTS:
            return True
        return False

    @staticmethod
    def _extract_tables(tables) -> list[dict]:
        result: list[dict] = []
        for table in tables:
            headers: list[str] = []
            rows: list[list[str]] = []
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if ri == 0:
                    headers = cells
                else:
                    rows.append(cells)
            result.append({"headers": headers, "rows": rows, "position": 0})
        return result
