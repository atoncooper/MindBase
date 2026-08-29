"""doc_extract.py — local document text extraction for 文件入库.

Usage:
    python doc_extract.py <path>

Prints exactly one JSON object on stdout:
    {"ok": true,  "title": "...", "text": "...", "ext": "pdf"}
    {"ok": false, "error": "..."}

Supported extensions: txt / md / markdown (plain text), pdf (pymupdf),
docx (python-docx), html / htm (stdlib HTMLParser, no third-party deps).

Anything on stderr is diagnostics only — the Rust caller reads stdout and
must never mix the two streams.
"""

import json
import os
import re
import sys

MAX_CHARS = 2_000_000  # hard cap; a 2M-char document is already unrealistic
TITLE_MAX_CHARS = 120


class ScriptError(Exception):
    """Fatal extraction problem; the message is user-facing."""


def _read_text_file(path):
    """Read a plain-text file, tolerating the common Windows encodings."""
    raw = open(path, "rb").read()
    for encoding in ("utf-8", "gbk", "utf-16"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _first_line_title(text, fallback):
    for line in text.splitlines():
        stripped = line.strip().lstrip("#").strip()
        if stripped:
            return stripped[:TITLE_MAX_CHARS]
    return fallback[:TITLE_MAX_CHARS]


def extract_txt(path):
    text = _read_text_file(path)
    title = _first_line_title(text, os.path.splitext(os.path.basename(path))[0])
    return title, text


def extract_pdf(path):
    try:
        # Prefer the modern package name: `import fitz` prints a deprecation
        # warning (to stdout!) on pymupdf >= 1.26, which would corrupt the
        # JSON protocol if it ran before the verdict line.
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz  # legacy shim, pymupdf < 1.26
        except ImportError as exc:
            raise ScriptError("缺少 PDF 解析依赖（pymupdf），请重启应用让其自动安装") from exc
    try:
        doc = fitz.open(path)
    except Exception as exc:
        raise ScriptError(f"无法打开 PDF：{exc}") from exc
    pages = []
    try:
        for page in doc:
            pages.append(page.get_text("text"))
        meta_title = ""
        try:
            meta_title = (doc.metadata or {}).get("title") or ""
        except Exception:
            pass
    finally:
        doc.close()
    text = "\n\n".join(part.strip() for part in pages if part and part.strip())
    if not text.strip():
        raise ScriptError("该 PDF 没有可提取的文本层（可能是扫描件），暂不支持 OCR")
    title = meta_title.strip() or _first_line_title(text, os.path.splitext(os.path.basename(path))[0])
    return title, text


def extract_docx(path):
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ScriptError("缺少 DOCX 解析依赖（python-docx），请重启应用让其自动安装") from exc
    try:
        document = docx.Document(path)
    except Exception as exc:
        raise ScriptError(f"无法打开 DOCX：{exc}") from exc
    blocks = [para.text for para in document.paragraphs if para.text and para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    text = "\n".join(blocks)
    if not text.strip():
        raise ScriptError("该 DOCX 没有可提取的正文")
    meta_title = ""
    try:
        meta_title = (document.core_properties.title or "").strip()
    except Exception:
        pass
    title = meta_title or _first_line_title(text, os.path.splitext(os.path.basename(path))[0])
    return title, text


def _collect_text(html):
    """Strip tags from an HTML fragment via stdlib (no deps)."""
    import html as html_mod
    from html.parser import HTMLParser

    class Collector(HTMLParser):
        """Collect visible text; script/style are dropped entirely."""

        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.chunks = []
            self._skipped = 0

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style"):
                self._skipped += 1
            elif tag in ("p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"):
                self.chunks.append("\n")

        def handle_endtag(self, tag):
            if tag in ("script", "style") and self._skipped > 0:
                self._skipped -= 1

        def handle_data(self, data):
            if self._skipped == 0 and data.strip():
                self.chunks.append(data)

    parser = Collector()
    try:
        parser.feed(html)
    except Exception:
        pass  # malformed HTML: keep whatever text was collected
    text = html_mod.unescape("".join(parser.chunks))
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text).strip()
    return text


# Minimum plausible article length below which the readability result is
# treated as a miss and the naive whole-page extraction is used instead.
READABILITY_MIN_CHARS = 80


def extract_html(path):
    import html as html_mod

    raw = _read_text_file(path)
    result = None
    # Preferred: readability main-content extraction (drops nav/footer/ads).
    # Requires readability-lxml, installed on demand by the app; silently
    # falls back to the naive whole-page strip when absent or when it fails.
    try:
        from readability import Document as _ReadabilityDoc

        doc = _ReadabilityDoc(raw)
        article_html = doc.summary(html_partial=True)
        text = _collect_text(article_html)
        if len(text) >= READABILITY_MIN_CHARS:
            result = (doc.short_title().strip(), text)
    except ImportError:
        pass
    except Exception:  # noqa: BLE001 — readability chokes on odd markup
        pass
    if result is None:
        text = _collect_text(raw)
        if not text:
            raise ScriptError("该 HTML 没有可提取的正文")
        match = re.search(r"<title[^>]*>(.*?)</title>", raw, re.IGNORECASE | re.DOTALL)
        title = ""
        if match:
            title = html_mod.unescape(match.group(1)).strip()
        result = (title, text)
    title, text = result
    if not text:
        raise ScriptError("该 HTML 没有可提取的正文")
    title = title or _first_line_title(text, os.path.splitext(os.path.basename(path))[0])
    return title, text


EXTRACTORS = {
    "txt": extract_txt,
    "md": extract_txt,
    "markdown": extract_txt,
    "pdf": extract_pdf,
    "docx": extract_docx,
    "html": extract_html,
    "htm": extract_html,
}


def main():
    # Windows pipes default to the console codepage (gbk on zh-CN), which
    # cannot encode many CJK-rare characters and kills the JSON print. The
    # Rust caller reads raw bytes as UTF-8, so always emit UTF-8.
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:  # noqa: BLE001 — older/odd stdio wrappers
        pass
    if len(sys.argv) != 2:
        print(json.dumps({"ok": False, "error": "参数错误：需要一个文件路径"}, ensure_ascii=False))
        return 2
    path = sys.argv[1]
    try:
        if not os.path.isfile(path):
            raise ScriptError(f"文件不存在：{path}")
        ext = os.path.splitext(path)[1].lower().lstrip(".")
        extractor = EXTRACTORS.get(ext)
        if extractor is None:
            raise ScriptError(f"不支持的文件类型：{ext or '(无扩展名)'}")
        title, text = extractor(path)
        text = text[:MAX_CHARS]
        if not text.strip():
            raise ScriptError("未提取到任何文本内容")
        print(json.dumps({"ok": True, "title": title, "text": text, "ext": ext}, ensure_ascii=False))
        return 0
    except ScriptError as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False))
        return 1
    except Exception as exc:  # noqa: BLE001 — the boundary must never crash raw
        print(json.dumps({"ok": False, "error": f"解析失败：{exc}"}, ensure_ascii=False))
        return 1


if __name__ == "__main__":
    sys.exit(main())
